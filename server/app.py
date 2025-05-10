import os
import json
import uuid
from datetime import datetime
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from dotenv import load_dotenv
from openai import OpenAI
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib import colors
import stripe
import time
# Import for DOCX generation
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "default-secret-key")
frontend_URL = os.getenv("FRONTEND_URL")

# Configure CORS
CORS(app, resources={r"/api/*": {
    "origins": [
        f"{frontend_URL}",
        "http://localhost:3000",
        "https://lexgenai.vercel.app",
        "https://lexgenai-git-main-mohamedaldahoul.vercel.app",
        "https://lexgenai.onrender.com"
    ],
    "methods": ["GET", "POST", "OPTIONS"],
    "allow_headers": ["Content-Type", "Authorization"],
    "expose_headers": ["Content-Type", "Authorization"],
    "supports_credentials": True,
    "max_age": 600
}})

# Configure rate limiting
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://",
)

# Configure Stripe
stripe.api_key = os.getenv("STRIPE_SECRET_KEY")

# Configure OpenAI
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

if not client.api_key:
    raise ValueError("OPENAI_API_KEY environment variable is not set")

# Document types and their descriptions
DOCUMENT_TYPES = {
    "nda": "Non-Disclosure Agreement (NDA)",
    "terms": "Website Terms of Service",
    "privacy": "Privacy Policy",
    "contract": "Freelance Contract",
    "employee": "Employment Agreement",
    "partnership": "Partnership Agreement"
}

# Ensure the downloads directory exists
DOWNLOAD_FOLDER = os.path.join(os.getcwd(), "static", "downloads")
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

# Add a Revisions folder
REVISIONS_FOLDER = os.path.join(os.getcwd(), "static", "revisions")
os.makedirs(REVISIONS_FOLDER, exist_ok=True)

# Configure test mode
TEST_MODE_ENABLED = os.getenv("ENABLE_TEST_MODE", "false").lower() == "true"

@app.route('/api/create-checkout-session', methods=['POST'])
@limiter.limit("10 per minute")
def create_checkout_session():
    try:
        data = request.get_json()
        form_data = data.get('formData', {})
        
        # Create a Stripe checkout session
        checkout_session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[
                {
                    'price_data': {
                        'currency': 'eur',
                        'product_data': {
                            'name': f'Legal Document: {DOCUMENT_TYPES.get(form_data.get("document_type", ""), "Custom Document")}',
                            'description': 'AI-generated legal document tailored to your business needs',
                        },
                        'unit_amount': 1900,  # £19.00 in cents
                    },
                    'quantity': 1,
                },
            ],
            mode='payment',
            success_url=f'{frontend_URL}/payment-processing?session_id={{CHECKOUT_SESSION_ID}}',
            cancel_url=f'{frontend_URL}/payment-processing?status=cancelled',
            metadata={
                'form_data': json.dumps(form_data)
            }
        )
        
        return jsonify({
            'sessionId': checkout_session.id
        })
    except Exception as e:
        app.logger.error(f"Stripe error: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/api/payment-success', methods=['GET'])
@limiter.limit("20 per minute")
def payment_success():
    session_id = request.args.get('session_id')
    
    try:
        session = stripe.checkout.Session.retrieve(session_id)
        form_data = json.loads(session.metadata.get('form_data', '{}'))
        
        start_time = time.time()
        timeout_limit = 25
        
        max_retries = 2
        retry_delay = 1
        
        for attempt in range(max_retries):
            try:
                if time.time() - start_time > timeout_limit:
                    return jsonify({
                        'status': 'processing',
                        'message': 'Your document is still being generated. Please wait a moment and try again.'
                    }), 202
                
                document_result = generate_document(form_data)
                
                if document_result.get('success'):
                    return jsonify(document_result)
                else:
                    error_msg = document_result.get('error', 'Unknown error')
                    app.logger.error(f"Document generation failed: {error_msg}")
                    
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                        retry_delay *= 2
                    else:
                        return jsonify({'error': f'Failed to generate document after {max_retries} attempts: {error_msg}'}), 500
            
            except Exception as doc_error:
                app.logger.error(f"Document generation exception (attempt {attempt+1}/{max_retries}): {str(doc_error)}")
                
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                    retry_delay *= 2
                else:
                    return jsonify({'error': f"Document generation failed after {max_retries} attempts: {str(doc_error)}"}), 500
            
    except stripe.error.StripeError as e:
        app.logger.error(f"Stripe error: {str(e)}")
        return jsonify({'error': f"Payment verification failed: {str(e)}"}), 400
    except Exception as e:
        app.logger.error(f"Payment success route error: {str(e)}")
        return jsonify({'error': f"An unexpected error occurred: {str(e)}"}), 500

def generate_document(form_data, generate_pdf=True, generate_docx=False):
    try:
        document_type = form_data.get('document_type')
        business_name = form_data.get('business_name')
        business_type = form_data.get('business_type')
        country = form_data.get('country')
        state_province = form_data.get('state_province', '')
        language = form_data.get('language')
        industry = form_data.get('industry')
        protection_level = form_data.get('protection_level', '2')
        
        clauses = []
        if form_data.get('clause_confidentiality'):
            clauses.append("Enhanced Confidentiality")
        if form_data.get('clause_arbitration'):
            clauses.append("Arbitration Provision")
        if form_data.get('clause_termination'):
            clauses.append("Advanced Termination Options")
        if form_data.get('clause_ip'):
            clauses.append("Intellectual Property Protection")
        
        additional_instructions = form_data.get('additional_instructions', '')
        
        # Include state/province in the prompt if provided
        location_detail = f"{country}"
        if state_province:
            location_detail = f"{state_province}, {country}"
        
        prompt = f"""Generate a professional {DOCUMENT_TYPES.get(document_type, 'legal document')} for {business_name}, a {business_type} in the {industry} industry, operating in {location_detail}.
Language document should be in {language}

Protection Level: {protection_level} out of 3

Special Clauses to Include: {', '.join(clauses) if clauses else 'None'}

Additional Instructions: {additional_instructions}

**Formatting Guidelines:**
- Use clear section headings in bold and all caps (e.g., **TERMS AND CONDITIONS**).
- Use proper indentation and line spacing for readability.
- Ensure signature fields are properly spaced and formatted as follows:

  **Signature:** ______________________  **Date:** _______________

- Use bullet points for lists where appropriate.
- Avoid overly dense paragraphs; break them up into short, digestible sections.
- Use legal language but ensure clarity for business professionals.

Format the document professionally with appropriate sections, headings, and legal language. Include all necessary legal provisions for this type of document in {location_detail}.
"""

        max_retries = 3
        retry_delay = 2
        
        for attempt in range(max_retries):
            try:
                response = client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a legal document generator that creates professional, legally-sound documents tailored to specific business needs and jurisdictions."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=4000,
                    temperature=0.7
                )
                document_text = response.choices[0].message.content
                break
            except Exception as e:
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                    retry_delay *= 2
                else:
                    raise Exception(f"Failed to generate document after {max_retries} attempts: {str(e)}")
        
        unique_id = uuid.uuid4().hex[:8]
        result = {
            'success': True,
            'preview': document_text
        }
        
        if generate_pdf:
            pdf_filename = f"{document_type}_{unique_id}.pdf"
            pdf_filepath = os.path.join(DOWNLOAD_FOLDER, pdf_filename)
            create_pdf(document_text, pdf_filepath, business_name, DOCUMENT_TYPES.get(document_type, "Legal Document"))
            result['pdf_filename'] = pdf_filename
        
        if generate_docx:
            docx_filename = f"{document_type}_{unique_id}.docx"
            docx_filepath = os.path.join(DOWNLOAD_FOLDER, docx_filename)
            create_docx(document_text, docx_filepath, business_name, DOCUMENT_TYPES.get(document_type, "Legal Document"))
            result['docx_filename'] = docx_filename
            
        return result
    
    except Exception as e:
        app.logger.error(f"Document generation error: {str(e)}")
        raise Exception(f"Failed to generate document: {str(e)}")

def create_pdf(text, filepath, business_name, document_type):
    doc = SimpleDocTemplate(filepath, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'Title',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=20,
        textColor=colors.navy,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'Normal',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        firstLineIndent=20,
        leading=14,
        spaceBefore=6,
        spaceAfter=6
    )
    
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Heading2'],
        fontSize=13,
        spaceAfter=10,
        spaceBefore=15,
        textColor=colors.navy,
        fontName='Helvetica-Bold',
        borderWidth=1,
        borderColor=colors.lightgrey,
        borderPadding=5,
        borderRadius=2
    )
    
    content = []
    
    content.append(Paragraph(f"{document_type.upper()}", title_style))
    content.append(Paragraph(f"For: {business_name}", title_style))
    content.append(Spacer(1, 20))
    
    date_style = ParagraphStyle(
        'Date',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_RIGHT,
        textColor=colors.darkgrey
    )
    content.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", date_style))
    content.append(Spacer(1, 20))
    
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            if para.strip().startswith('#'):
                header_text = para.replace('#', '').strip()
                content.append(Paragraph(header_text, header_style))
            elif para.strip().isupper() and len(para.strip()) > 3:
                content.append(Paragraph(para.strip(), header_style))
            elif para.strip().startswith(('•', '-', '*')):
                bullet_style = ParagraphStyle(
                    'Bullet',
                    parent=normal_style,
                    leftIndent=30,
                    firstLineIndent=0,
                    spaceBefore=3,
                    spaceAfter=3
                )
                content.append(Paragraph(para.strip(), bullet_style))
            elif "signature" in para.lower() or "sign" in para.lower() or "date:" in para.lower():
                sig_style = ParagraphStyle(
                    'Signature',
                    parent=normal_style,
                    spaceBefore=15,
                    spaceAfter=15
                )
                content.append(Paragraph(para, sig_style))
            else:
                content.append(Paragraph(para, normal_style))
            
            if para.strip().startswith('#') or para.strip().isupper():
                content.append(Spacer(1, 10))
            else:
                content.append(Spacer(1, 6))
    
    doc.build(content)

def create_docx(text, filepath, business_name, document_type):
    doc = Document()
    
    # Set document margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run(document_type.upper())
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add business name
    business = doc.add_paragraph()
    business_run = business.add_run(f"For: {business_name}")
    business_run.bold = True
    business_run.font.size = Pt(14)
    business.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add document text
    paragraphs = text.split("\n")
    for para in paragraphs:
        if not para.strip():
            continue
            
        # Determine paragraph style
        if para.strip().startswith('#') or (para.strip().isupper() and len(para.strip()) > 3):
            # This is a heading
            header_text = para.replace('#', '').strip()
            header = doc.add_paragraph()
            header_run = header.add_run(header_text if '#' in para else para)
            header_run.bold = True
            header_run.font.size = Pt(14)
            header.style = 'Heading 2'
        
        elif para.strip().startswith(('•', '-', '*')):
            # This is a bullet point
            p = doc.add_paragraph(para.strip().lstrip('•-* '), style='List Bullet')
            
        elif "signature" in para.lower() or "sign" in para.lower() or "date:" in para.lower():
            # This is a signature line
            p = doc.add_paragraph()
            p.add_run(para).bold = True
            p.space_after = Pt(20)
            
        else:
            # Regular paragraph
            p = doc.add_paragraph(para)
    
    doc.save(filepath)

@app.route('/api/download/<filename>')
# @limiter.limit("10 per minute")
def download_file(filename):
    return send_from_directory(DOWNLOAD_FOLDER, filename, as_attachment=True)

@app.route('/api/health', methods=['GET'])
def health_check():
    try:
        stripe.Account.retrieve()
        openai_status = "ok"
        try:
            client.models.list()
        except Exception as e:
            openai_status = f"error: {str(e)}"
        
        return jsonify({
            'status': 'healthy',
            'timestamp': datetime.now().isoformat(),
            'stripe': 'ok',
            'openai': openai_status
        })
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 500

def is_localhost():
    remote_addr = request.remote_addr
    return remote_addr == "127.0.0.1" or remote_addr == "localhost" or remote_addr.startswith("192.168.") or remote_addr.startswith("10.")

@app.route('/api/preview-document', methods=['POST'])
def preview_document():
    if not is_localhost():
        return jsonify({
            "error": "This endpoint is only available in development mode",
            "status": "error"
        }), 403

    try:
        data = request.json
        
        # Add state/province to location if available
        location = data['country']
        if data.get('state_province'):
            location = f"{data['state_province']}, {data['country']}"
        
        # Create a prompt based on the form data
        prompt = f"""Create a {data['document_type']} for {data['business_name']}, a {data['business_type']} in {location}.
        Industry: {data['industry']}
        Protection Level: {data['protection_level']}
        Special Clauses:
        - Confidentiality: {'Yes' if data.get('clause_confidentiality') else 'No'}
        - Arbitration: {'Yes' if data.get('clause_arbitration') else 'No'}
        - Termination: {'Yes' if data.get('clause_termination') else 'No'}
        - IP Protection: {'Yes' if data.get('clause_ip') else 'No'}
        
        Additional Instructions: {data.get('additional_instructions', 'None')}
        
        Please generate a professional legal document based on these requirements."""

        # Call OpenAI API with new syntax
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a legal document generation assistant. Create professional, well-structured legal documents based on the provided requirements."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.7
        )

        # Extract the generated document with new syntax
        generated_text = response.choices[0].message.content

        return jsonify({
            "preview": generated_text,
            "status": "success"
        })

    except Exception as e:
        return jsonify({
            "error": str(e),
            "status": "error"
        }), 500

@app.route('/api/generate-test-document', methods=['POST'])
def generate_test_document():
    if not is_localhost():
        return jsonify({
            "error": "This endpoint is only available in development mode",
            "status": "error"
        }), 403
        
    try:
        data = request.json
        document_result = generate_document(data, generate_pdf=True, generate_docx=False)
        
        if document_result.get('success'):
            return jsonify(document_result)
        else:
            return jsonify({
                "error": "Failed to generate document",
                "status": "error"
            }), 500

    except Exception as e:
        return jsonify({
            "error": str(e),
            "status": "error"
        }), 500

@app.route('/api/document-details', methods=['GET'])
def document_details():
    session_id = request.args.get('session_id')
    if not session_id:
        return jsonify({'error': 'No session_id provided'}), 400

    try:
        session = stripe.checkout.Session.retrieve(session_id)
        form_data = json.loads(session.metadata.get('form_data', '{}'))
        document_result = generate_document(form_data, generate_pdf=False, generate_docx=False)
        if document_result.get('success'):
            return jsonify({
                'preview': document_result.get('preview')
            })
        else:
            return jsonify({'error': 'Document not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-pdf', methods=['GET'])
def generate_pdf_on_demand():
    session_id = request.args.get('session_id')
    if not session_id:
        return jsonify({'error': 'No session_id provided'}), 400

    try:
        session = stripe.checkout.Session.retrieve(session_id)
        form_data = json.loads(session.metadata.get('form_data', '{}'))
        document_result = generate_document(form_data, generate_pdf=True, generate_docx=False)
        if document_result.get('success'):
            filename = document_result.get('pdf_filename')
            return send_from_directory(DOWNLOAD_FOLDER, filename, as_attachment=True)
        else:
            return jsonify({'error': 'Failed to generate PDF'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/generate-docx', methods=['GET'])
def generate_docx_on_demand():
    session_id = request.args.get('session_id')
    if not session_id:
        return jsonify({'error': 'No session_id provided'}), 400

    try:
        session = stripe.checkout.Session.retrieve(session_id)
        form_data = json.loads(session.metadata.get('form_data', '{}'))
        document_result = generate_document(form_data, generate_pdf=False, generate_docx=True)
        if document_result.get('success'):
            filename = document_result.get('docx_filename')
            return send_from_directory(DOWNLOAD_FOLDER, filename, as_attachment=True)
        else:
            return jsonify({'error': 'Failed to generate DOCX'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/document-feedback', methods=['POST'])
def submit_document_feedback():
    try:
        data = request.get_json()
        session_id = data.get('sessionId')
        comment = data.get('comment')
        
        if not session_id or not comment:
            return jsonify({'error': 'Both sessionId and comment are required'}), 400
            
        # Retrieve the session to verify it exists
        try:
            session = stripe.checkout.Session.retrieve(session_id)
            form_data = json.loads(session.metadata.get('form_data', '{}'))
        except Exception as e:
            return jsonify({'error': f'Invalid session: {str(e)}'}), 400
            
        # Create a revision record
        revision_id = uuid.uuid4().hex
        feedback_file = os.path.join(DOWNLOAD_FOLDER, f"feedback_{revision_id}.json")
        
        revision_data = {
            'session_id': session_id,
            'revision_id': revision_id,
            'comment': comment,
            'timestamp': datetime.now().isoformat(),
            'status': 'pending',
            'form_data': form_data
        }
        
        with open(feedback_file, 'w') as f:
            json.dump(revision_data, f, indent=2)
            
        # Schedule or trigger document update process
        # For this implementation, we'll generate the updated document immediately
        try:
            updated_document = generate_revised_document(revision_data)
            
            # Update the revision status
            revision_data['status'] = 'completed'
            revision_data['completed_at'] = datetime.now().isoformat()
            
            with open(feedback_file, 'w') as f:
                json.dump(revision_data, f, indent=2)
                
            return jsonify({
                'success': True,
                'message': 'Your document has been updated successfully',
                'revision_id': revision_id
            })
            
        except Exception as update_error:
            app.logger.error(f"Error updating document: {str(update_error)}")
            # Mark the revision as failed
            revision_data['status'] = 'failed'
            revision_data['error'] = str(update_error)
            
            with open(feedback_file, 'w') as f:
                json.dump(revision_data, f, indent=2)
                
            return jsonify({
                'success': False,
                'message': 'We received your feedback but could not generate an updated document. Our team will review it.',
                'revision_id': revision_id
            }), 500
            
    except Exception as e:
        app.logger.error(f"Error submitting feedback: {str(e)}")
        return jsonify({'error': str(e)}), 500

def generate_revised_document(revision_data):
    """Generate an updated document based on user feedback"""
    session_id = revision_data['session_id']
    revision_id = revision_data['revision_id']
    comment = revision_data['comment']
    form_data = revision_data['form_data']
    
    # Get the original document text by generating it again
    # In a production system, you might want to store and retrieve the original text
    original_result = generate_document(form_data, generate_pdf=False, generate_docx=False)
    original_text = original_result.get('preview', '')
    
    # Create a prompt for updating the document
    prompt = f"""
I have a legal document that needs to be updated based on user feedback. 

Here is the original document:
```
{original_text}
```

The user has requested the following changes:
```
{comment}
```

Please provide the complete updated document with the requested changes incorporated. 
Return only the revised document text, properly formatted with all original sections and with the requested changes applied.
"""

    # Generate the updated document using OpenAI
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a legal document revision assistant. Update documents precisely according to user feedback while maintaining their professional structure and format."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4000,
        temperature=0.7
    )
    
    updated_text = response.choices[0].message.content
    
    # Save the updated document in various formats
    document_type = form_data.get('document_type', 'document')
    business_name = form_data.get('business_name', 'Business')
    
    # Generate PDF
    pdf_filename = f"{document_type}_rev_{revision_id}.pdf"
    pdf_filepath = os.path.join(REVISIONS_FOLDER, pdf_filename)
    create_pdf(updated_text, pdf_filepath, business_name, DOCUMENT_TYPES.get(document_type, "Legal Document"))
    
    # Generate DOCX
    docx_filename = f"{document_type}_rev_{revision_id}.docx"
    docx_filepath = os.path.join(REVISIONS_FOLDER, docx_filename)
    create_docx(updated_text, docx_filepath, business_name, DOCUMENT_TYPES.get(document_type, "Legal Document"))
    
    # Store revision info
    revision_info = {
        'session_id': session_id,
        'revision_id': revision_id,
        'original_text': original_text,
        'updated_text': updated_text,
        'pdf_filename': pdf_filename,
        'docx_filename': docx_filename,
        'timestamp': datetime.now().isoformat()
    }
    
    revision_info_file = os.path.join(REVISIONS_FOLDER, f"revision_info_{revision_id}.json")
    with open(revision_info_file, 'w') as f:
        json.dump(revision_info, f, indent=2)
    
    return revision_info

@app.route('/api/document-revisions/<session_id>', methods=['GET'])
def get_document_revisions(session_id):
    """Get all revisions for a document session"""
    if not session_id:
        return jsonify({'error': 'Session ID is required'}), 400
        
    try:
        # Verify the session exists
        stripe.checkout.Session.retrieve(session_id)
        
        # Find all revisions for this session
        revisions = []
        
        # Search in the downloads folder for feedback files
        for file in os.listdir(DOWNLOAD_FOLDER):
            if file.startswith("feedback_") and file.endswith(".json"):
                try:
                    with open(os.path.join(DOWNLOAD_FOLDER, file), 'r') as f:
                        feedback_data = json.load(f)
                        if feedback_data.get('session_id') == session_id:
                            revisions.append({
                                'revision_id': feedback_data.get('revision_id'),
                                'status': feedback_data.get('status'),
                                'timestamp': feedback_data.get('timestamp'),
                                'comment': feedback_data.get('comment')
                            })
                except Exception as e:
                    app.logger.error(f"Error reading feedback file {file}: {str(e)}")
        
        return jsonify({
            'session_id': session_id,
            'revisions': sorted(revisions, key=lambda x: x.get('timestamp', ''), reverse=True)
        })
        
    except Exception as e:
        app.logger.error(f"Error retrieving revisions: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/revised-document/<revision_id>', methods=['GET'])
def get_revised_document(revision_id):
    """Get the revised document preview"""
    if not revision_id:
        return jsonify({'error': 'Revision ID is required'}), 400
        
    try:
        # Look for the revision info file
        revision_info_file = os.path.join(REVISIONS_FOLDER, f"revision_info_{revision_id}.json")
        
        if not os.path.exists(revision_info_file):
            return jsonify({'error': 'Revision not found'}), 404
            
        with open(revision_info_file, 'r') as f:
            revision_info = json.load(f)
            
        return jsonify({
            'revision_id': revision_id,
            'preview': revision_info.get('updated_text')
        })
        
    except Exception as e:
        app.logger.error(f"Error retrieving revised document: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/download-revision/<format>/<revision_id>', methods=['GET'])
def download_revised_document(format, revision_id):
    """Download a revised document in the specified format"""
    if not revision_id or format not in ['pdf', 'docx']:
        return jsonify({'error': 'Valid revision ID and format (pdf/docx) are required'}), 400
        
    try:
        # Look for the revision info file
        revision_info_file = os.path.join(REVISIONS_FOLDER, f"revision_info_{revision_id}.json")
        
        if not os.path.exists(revision_info_file):
            return jsonify({'error': 'Revision not found'}), 404
            
        with open(revision_info_file, 'r') as f:
            revision_info = json.load(f)
            
        filename = revision_info.get(f'{format}_filename')
        
        if not filename:
            return jsonify({'error': f'No {format.upper()} file found for this revision'}), 404
            
        return send_from_directory(REVISIONS_FOLDER, filename, as_attachment=True)
        
    except Exception as e:
        app.logger.error(f"Error downloading revised document: {str(e)}")
        return jsonify({'error': str(e)}), 500

# Add OPTIONS handler for all routes
@app.after_request
def after_request(response):
    response.headers.add('Access-Control-Allow-Origin', 'http://localhost:3000')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,PUT,POST,DELETE,OPTIONS')
    return response

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port) 